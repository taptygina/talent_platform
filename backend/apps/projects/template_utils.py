import re
from io import BytesIO
from typing import Any

from django.utils.text import slugify

from apps.projects.models import Project, ProjectTemplate, ProjectTemplateSection
from apps.users.models import User


PLACEHOLDER_RE = re.compile(r"\{\{\s*([A-ZА-Я_]+)\s*\}\}")
HEADING_PREFIX_RE = re.compile(r"^\s*\d+(?:\.\d+)*[\)\.]?\s+")
NUMBERED_HEADING_RE = re.compile(r"^\s*(\d+(?:\.\d+){0,4})\.?\s+(.+?)\s*$")
TOC_DOTS_RE = re.compile(r"\.{2,}\s*\d+\s*$")
TOC_PAGE_RE = re.compile(r"^.{3,}\s+\d{1,3}\s*$")
TRAILING_PAGE_RE = re.compile(r"^(.*\D)\s+(\d{1,3})$")
APPENDIX_HEADING_RE = re.compile(r"^\s*\u043f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u0435\s+[\u0430-\u044f\u0451a-z0-9]+\s*\.?\s*$", re.IGNORECASE)
HEADING_STYLE_LEVEL_RE = re.compile(r"^heading\s*(\d+)$", re.IGNORECASE)

SPECIAL_SECTION_TITLES = {
    "введение",
    "заключение",
    "список использованных источников",
    "список литературы",
    "приложение",
}

BUILDER_VARIABLES = {
    "НАЗВАНИЕ_ПРОЕКТА": "Название проекта",
    "ЦЕЛЬ_ПРОЕКТА": "Цель проекта",
    "ОПИСАНИЕ_ПРОЕКТА": "Описание проекта",
    "ФИО_РУКОВОДИТЕЛЯ": "ФИО руководителя",
    "ФИО_СТУДЕНТА": "ФИО студента",
    "ГРУППА_СТУДЕНТА": "Группа студента",
    "НАЗВАНИЕ_ЭТАПА": "Название этапа",
    "ЗАДАНИЕ_ЭТАПА": "Задание этапа",
    "ОТЧЕТ_ПО_ЭТАПУ": "Отчет студента",
    "PROJECT_TITLE": "Название проекта",
    "PROJECT_GOAL": "Цель проекта",
    "PROJECT_DESCRIPTION": "Описание проекта",
    "SUPERVISOR_FULL_NAME": "ФИО руководителя",
    "STUDENT_FULL_NAME": "ФИО студента",
    "STUDENT_GROUP": "Группа студента",
    "STAGE_TITLE": "Название этапа",
    "STAGE_TASK": "Задание этапа",
    "STAGE_REPORT": "Отчет студента",
}

BUILDER_REPEAT_SOURCES = {"participants", "stages"}
BUILDER_CONDITION_KEYS = {
    "has_goal",
    "has_description",
    "has_participants",
    "has_stages",
}


def _full_name(user: User | None) -> str:
    if not user:
        return ""
    if getattr(user, "full_name", ""):
        return user.full_name
    parts = [getattr(user, "last_name", ""), getattr(user, "first_name", ""), getattr(user, "middle_name", "")]
    fallback = " ".join([part for part in parts if part]).strip()
    return fallback or getattr(user, "username", "")


def render_placeholders(
    text: str,
    *,
    student: User | None = None,
    supervisor: User | None = None,
    project: Project | None = None,
) -> str:
    source = text or ""
    mapping = {
        "STUDENT_FULL_NAME": _full_name(student),
        "STUDENT_GROUP": getattr(student, "group_name", "") if student else "",
        "TEACHER_FULL_NAME": _full_name(supervisor),
        "SUPERVISOR_FULL_NAME": _full_name(supervisor),
        "PROJECT_TITLE": getattr(project, "title", "") if project else "",
        "ФИО_СТУДЕНТА": _full_name(student),
        "ГРУППА_СТУДЕНТА": getattr(student, "group_name", "") if student else "",
        "ФИО_ПРЕПОДАВАТЕЛЯ": _full_name(supervisor),
        "ФИО_РУКОВОДИТЕЛЯ": _full_name(supervisor),
        "НАЗВАНИЕ_ПРОЕКТА": getattr(project, "title", "") if project else "",
    }

    def replace_match(match: re.Match[str]) -> str:
        key = (match.group(1) or "").strip().upper()
        return mapping.get(key, match.group(0))

    return PLACEHOLDER_RE.sub(replace_match, source)


def default_builder_schema() -> dict[str, Any]:
    return {
        "version": 1,
        "blocks": [
            {"id": "heading-main", "type": "heading", "level": 1, "text": "{{НАЗВАНИЕ_ПРОЕКТА}}"},
            {"id": "paragraph-goal", "type": "paragraph", "text": "Цель проекта: {{ЦЕЛЬ_ПРОЕКТА}}"},
            {
                "id": "repeat-stages",
                "type": "repeat",
                "source": "stages",
                "children": [
                    {"id": "stage-heading", "type": "heading", "level": 2, "text": "{{НАЗВАНИЕ_ЭТАПА}}"},
                    {"id": "stage-task", "type": "paragraph", "text": "{{ЗАДАНИЕ_ЭТАПА}}"},
                ],
            },
        ],
    }


def _clean_block(block: dict[str, Any], depth: int = 0) -> dict[str, Any] | None:
    if depth > 6 or not isinstance(block, dict):
        return None
    block_type = str(block.get("type") or "").strip()
    if block_type not in {"heading", "paragraph", "variable", "condition", "repeat", "page_break"}:
        return None
    clean = {"id": str(block.get("id") or "").strip() or f"block-{depth}"}
    clean["type"] = block_type
    if block_type in {"heading", "paragraph"}:
        clean["text"] = str(block.get("text") or "")
    if block_type == "heading":
        try:
            clean["level"] = max(1, min(3, int(block.get("level") or 1)))
        except Exception:
            clean["level"] = 1
    if block_type == "variable":
        key = str(block.get("key") or "НАЗВАНИЕ_ПРОЕКТА").strip().upper()
        clean["key"] = key if key in BUILDER_VARIABLES else "НАЗВАНИЕ_ПРОЕКТА"
    if block_type == "condition":
        key = str(block.get("key") or "has_goal").strip()
        clean["key"] = key if key in BUILDER_CONDITION_KEYS else "has_goal"
        clean["children"] = [item for child in block.get("children") or [] if (item := _clean_block(child, depth + 1))]
    if block_type == "repeat":
        source = str(block.get("source") or "stages").strip()
        clean["source"] = source if source in BUILDER_REPEAT_SOURCES else "stages"
        clean["children"] = [item for child in block.get("children") or [] if (item := _clean_block(child, depth + 1))]
    return clean


def normalize_builder_schema(value: Any) -> dict[str, Any]:
    raw = value if isinstance(value, dict) else {}
    blocks = [item for child in raw.get("blocks") or [] if (item := _clean_block(child))]
    return {"version": 1, "blocks": blocks}


def _builder_base_context(project: Project | None = None, student: User | None = None, stage: Any = None) -> dict[str, str]:
    supervisor = getattr(project, "supervisor", None) if project else None
    context = {
        "PROJECT_TITLE": getattr(project, "title", "") if project else "Цифровой паспорт компетенций",
        "PROJECT_GOAL": getattr(project, "goal", "") if project else "Сократить время подбора участников под проектные задачи.",
        "PROJECT_DESCRIPTION": getattr(project, "description", "") if project else "Единый профиль участника с навыками и проектным опытом.",
        "SUPERVISOR_FULL_NAME": _full_name(supervisor) if supervisor else "Петров Иван Алексеевич",
        "STUDENT_FULL_NAME": _full_name(student) if student else "Соколова Анна Ильинична",
        "STUDENT_GROUP": getattr(student, "group_name", "") if student else "ИС-222б",
        "STAGE_TITLE": getattr(stage, "title", "") if stage else "Этап 1. Постановка задачи",
        "STAGE_TASK": getattr(stage, "task_text", "") if stage else "Собрать требования и согласовать критерии результата.",
        "STAGE_REPORT": getattr(stage, "student_report", "") if stage else "Материалы этапа подготовлены и загружены.",
    }
    context.update({
        "НАЗВАНИЕ_ПРОЕКТА": context["PROJECT_TITLE"],
        "ЦЕЛЬ_ПРОЕКТА": context["PROJECT_GOAL"],
        "ОПИСАНИЕ_ПРОЕКТА": context["PROJECT_DESCRIPTION"],
        "ФИО_РУКОВОДИТЕЛЯ": context["SUPERVISOR_FULL_NAME"],
        "ФИО_СТУДЕНТА": context["STUDENT_FULL_NAME"],
        "ГРУППА_СТУДЕНТА": context["STUDENT_GROUP"],
        "НАЗВАНИЕ_ЭТАПА": context["STAGE_TITLE"],
        "ЗАДАНИЕ_ЭТАПА": context["STAGE_TASK"],
        "ОТЧЕТ_ПО_ЭТАПУ": context["STAGE_REPORT"],
    })
    return context


def _render_builder_text(text: str, context: dict[str, str]) -> str:
    def replace(match: re.Match[str]) -> str:
        return context.get((match.group(1) or "").strip().upper(), match.group(0))

    return PLACEHOLDER_RE.sub(replace, text or "")


def _condition_value(key: str, *, project: Project | None, participants: list[User], stages: list[Any]) -> bool:
    return {
        "has_goal": bool(getattr(project, "goal", "") if project else True),
        "has_description": bool(getattr(project, "description", "") if project else True),
        "has_participants": bool(participants),
        "has_stages": bool(stages),
    }.get(key, False)


def render_builder_blocks(schema: dict[str, Any], *, project: Project | None = None) -> list[dict[str, Any]]:
    normalized = normalize_builder_schema(schema)
    participants = list(project.participants.all()) if project else []
    stages = list(project.stages.order_by("order", "id")) if project else []
    if not participants:
        participants = [None]
    if not stages:
        stages = [None]

    rendered: list[dict[str, Any]] = []

    def walk(blocks: list[dict[str, Any]], *, student: User | None = None, stage: Any = None):
        for block in blocks:
            block_type = block["type"]
            if block_type in {"heading", "paragraph"}:
                rendered.append({
                    "type": block_type,
                    "level": block.get("level", 1),
                    "text": _render_builder_text(block.get("text", ""), _builder_base_context(project, student, stage)),
                })
            elif block_type == "variable":
                rendered.append({
                    "type": "paragraph",
                    "level": 1,
                    "text": _builder_base_context(project, student, stage).get(block.get("key", ""), ""),
                })
            elif block_type == "page_break":
                rendered.append({"type": "page_break"})
            elif block_type == "condition":
                if _condition_value(block.get("key", ""), project=project, participants=participants if participants != [None] else [], stages=stages if stages != [None] else []):
                    walk(block.get("children", []), student=student, stage=stage)
            elif block_type == "repeat":
                if block.get("source") == "participants":
                    for item in participants:
                        walk(block.get("children", []), student=item, stage=stage)
                else:
                    for item in stages:
                        walk(block.get("children", []), student=student, stage=item)

    walk(normalized["blocks"])
    return rendered


def _iter_doc_strings(doc: Any):
    for paragraph in doc.paragraphs:
        yield paragraph.text or ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph.text or ""


def _safe_pt(value: Any) -> float | None:
    if value is None:
        return None
    try:
        return round(float(value.pt), 2)
    except Exception:
        return None


def _safe_cm(value: Any) -> float | None:
    if value is None:
        return None
    try:
        return round(float(value.cm), 2)
    except Exception:
        return None


def _pick_heading_paragraph(doc: Any):
    for paragraph in doc.paragraphs:
        raw = (paragraph.text or "").strip()
        if not raw:
            continue
        if _is_heading_paragraph(paragraph):
            return paragraph
        style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").strip().lower()
        if style_name.startswith("heading"):
            return paragraph
        if _is_numbered_heading_text(raw) or _is_special_heading_text(raw):
            return paragraph
    return None


def _pick_content_paragraph(doc: Any):
    for paragraph in doc.paragraphs:
        raw = (paragraph.text or "").strip()
        if not raw:
            continue
        style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").strip().lower()
        if style_name.startswith("toc") or style_name.startswith("heading") or style_name.startswith("title"):
            continue
        return paragraph
    for paragraph in doc.paragraphs:
        if (paragraph.text or "").strip():
            return paragraph
    return None


def _extract_line_spacing_config(paragraph_format: Any) -> tuple[float | None, str | None, int | None]:
    if paragraph_format is None:
        return None, None, None

    line_spacing = None
    line_spacing_kind = None
    line_spacing_rule = None

    rule = getattr(paragraph_format, "line_spacing_rule", None)
    if rule is not None:
        try:
            line_spacing_rule = int(rule)
        except Exception:
            line_spacing_rule = None

    raw_spacing = getattr(paragraph_format, "line_spacing", None)
    if raw_spacing is None:
        return None, None, line_spacing_rule

    if hasattr(raw_spacing, "pt"):
        pt_value = _safe_pt(raw_spacing)
        if pt_value is not None:
            line_spacing = pt_value
            line_spacing_kind = "pt"
            return line_spacing, line_spacing_kind, line_spacing_rule

    try:
        line_spacing = float(raw_spacing)
        line_spacing_kind = "multiple"
    except Exception:
        line_spacing = None
        line_spacing_kind = None

    return line_spacing, line_spacing_kind, line_spacing_rule


def _pick_non_empty(*values):
    for value in values:
        if value is not None:
            return value
    return None


def extract_template_format_profile_from_doc(doc: Any) -> dict[str, Any]:
    profile: dict[str, Any] = {
        "defaults": {
            "font_family": "",
            "font_size_pt": None,
        },
        "paragraph": {
            "line_spacing": None,
            "line_spacing_kind": None,
            "line_spacing_rule": None,
            "space_before_pt": None,
            "space_after_pt": None,
            "first_line_indent_cm": None,
            "left_indent_cm": None,
            "right_indent_cm": None,
            "alignment": None,
            "style_name": "",
        },
        "heading_paragraph": {
            "line_spacing": None,
            "line_spacing_kind": None,
            "line_spacing_rule": None,
            "space_before_pt": None,
            "space_after_pt": None,
            "first_line_indent_cm": None,
            "left_indent_cm": None,
            "right_indent_cm": None,
            "alignment": None,
            "style_name": "",
        },
        "page": {
            "margin_top_cm": None,
            "margin_right_cm": None,
            "margin_bottom_cm": None,
            "margin_left_cm": None,
            "header_distance_cm": None,
            "footer_distance_cm": None,
        },
        "header_footer": {
            "header_text": "",
            "footer_text": "",
        },
        "placeholders": [],
    }

    normal_style = doc.styles["Normal"]
    normal_font = normal_style.font

    content_paragraph = _pick_content_paragraph(doc)
    heading_paragraph = _pick_heading_paragraph(doc)

    content_pf = content_paragraph.paragraph_format if content_paragraph is not None else None
    content_style_pf = getattr(getattr(content_paragraph, "style", None), "paragraph_format", None) if content_paragraph is not None else None

    heading_pf = heading_paragraph.paragraph_format if heading_paragraph is not None else None
    heading_style_pf = getattr(getattr(heading_paragraph, "style", None), "paragraph_format", None) if heading_paragraph is not None else None

    run_font_name = None
    run_font_size_pt = None
    if content_paragraph is not None:
        for run in content_paragraph.runs:
            if not (run.text or "").strip():
                continue
            if getattr(run.font, "name", None):
                run_font_name = run.font.name
            if getattr(run.font, "size", None):
                run_font_size_pt = _safe_pt(run.font.size)
            if run_font_name or run_font_size_pt is not None:
                break

    profile["defaults"]["font_family"] = _pick_non_empty(run_font_name, getattr(normal_font, "name", None), "")
    profile["defaults"]["font_size_pt"] = _pick_non_empty(run_font_size_pt, _safe_pt(getattr(normal_font, "size", None)))

    normal_pf = normal_style.paragraph_format
    sample_pf = content_pf if content_pf is not None else (content_style_pf if content_style_pf is not None else normal_pf)

    ls, ls_kind, ls_rule = _extract_line_spacing_config(sample_pf)
    if ls is None and content_style_pf is not None:
        ls, ls_kind, ls_rule = _extract_line_spacing_config(content_style_pf)
    if ls is None and normal_pf is not None:
        ls, ls_kind, ls_rule = _extract_line_spacing_config(normal_pf)

    profile["paragraph"]["line_spacing"] = ls
    profile["paragraph"]["line_spacing_kind"] = ls_kind
    profile["paragraph"]["line_spacing_rule"] = ls_rule

    profile["paragraph"]["space_before_pt"] = _pick_non_empty(
        _safe_pt(getattr(sample_pf, "space_before", None)) if sample_pf is not None else None,
        _safe_pt(getattr(content_style_pf, "space_before", None)) if content_style_pf is not None else None,
        _safe_pt(getattr(normal_pf, "space_before", None)) if normal_pf is not None else None,
    )
    profile["paragraph"]["space_after_pt"] = _pick_non_empty(
        _safe_pt(getattr(sample_pf, "space_after", None)) if sample_pf is not None else None,
        _safe_pt(getattr(content_style_pf, "space_after", None)) if content_style_pf is not None else None,
        _safe_pt(getattr(normal_pf, "space_after", None)) if normal_pf is not None else None,
    )
    profile["paragraph"]["first_line_indent_cm"] = _pick_non_empty(
        _safe_cm(getattr(sample_pf, "first_line_indent", None)) if sample_pf is not None else None,
        _safe_cm(getattr(content_style_pf, "first_line_indent", None)) if content_style_pf is not None else None,
        _safe_cm(getattr(normal_pf, "first_line_indent", None)) if normal_pf is not None else None,
    )
    profile["paragraph"]["left_indent_cm"] = _pick_non_empty(
        _safe_cm(getattr(sample_pf, "left_indent", None)) if sample_pf is not None else None,
        _safe_cm(getattr(content_style_pf, "left_indent", None)) if content_style_pf is not None else None,
        _safe_cm(getattr(normal_pf, "left_indent", None)) if normal_pf is not None else None,
    )
    profile["paragraph"]["right_indent_cm"] = _pick_non_empty(
        _safe_cm(getattr(sample_pf, "right_indent", None)) if sample_pf is not None else None,
        _safe_cm(getattr(content_style_pf, "right_indent", None)) if content_style_pf is not None else None,
        _safe_cm(getattr(normal_pf, "right_indent", None)) if normal_pf is not None else None,
    )

    align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
    align_value = getattr(content_paragraph, "alignment", None) if content_paragraph is not None else None
    if align_value is None and sample_pf is not None:
        align_value = getattr(sample_pf, "alignment", None)
    if align_value is None and content_style_pf is not None:
        align_value = getattr(content_style_pf, "alignment", None)
    if align_value is None and normal_pf is not None:
        align_value = getattr(normal_pf, "alignment", None)
    if align_value is not None:
        try:
            profile["paragraph"]["alignment"] = align_map.get(int(align_value), "left")
        except Exception:
            profile["paragraph"]["alignment"] = "left"

    profile["paragraph"]["style_name"] = (getattr(getattr(content_paragraph, "style", None), "name", "") or "").strip() if content_paragraph is not None else ""

    # Профиль форматирования заголовков
    heading_base_pf = heading_pf if heading_pf is not None else (heading_style_pf if heading_style_pf is not None else None)
    h_ls, h_ls_kind, h_ls_rule = _extract_line_spacing_config(heading_base_pf)
    profile["heading_paragraph"]["line_spacing"] = h_ls
    profile["heading_paragraph"]["line_spacing_kind"] = h_ls_kind
    profile["heading_paragraph"]["line_spacing_rule"] = h_ls_rule
    profile["heading_paragraph"]["space_before_pt"] = _safe_pt(getattr(heading_base_pf, "space_before", None)) if heading_base_pf is not None else None
    profile["heading_paragraph"]["space_after_pt"] = _safe_pt(getattr(heading_base_pf, "space_after", None)) if heading_base_pf is not None else None
    profile["heading_paragraph"]["first_line_indent_cm"] = _safe_cm(getattr(heading_base_pf, "first_line_indent", None)) if heading_base_pf is not None else None
    profile["heading_paragraph"]["left_indent_cm"] = _safe_cm(getattr(heading_base_pf, "left_indent", None)) if heading_base_pf is not None else None
    profile["heading_paragraph"]["right_indent_cm"] = _safe_cm(getattr(heading_base_pf, "right_indent", None)) if heading_base_pf is not None else None
    h_align_value = getattr(heading_paragraph, "alignment", None) if heading_paragraph is not None else None
    if h_align_value is None and heading_base_pf is not None:
        h_align_value = getattr(heading_base_pf, "alignment", None)
    if h_align_value is not None:
        try:
            profile["heading_paragraph"]["alignment"] = align_map.get(int(h_align_value), "left")
        except Exception:
            profile["heading_paragraph"]["alignment"] = "left"
    profile["heading_paragraph"]["style_name"] = (getattr(getattr(heading_paragraph, "style", None), "name", "") or "").strip() if heading_paragraph is not None else ""

    if doc.sections:
        section = doc.sections[0]
        profile["page"]["margin_top_cm"] = _safe_cm(section.top_margin)
        profile["page"]["margin_right_cm"] = _safe_cm(section.right_margin)
        profile["page"]["margin_bottom_cm"] = _safe_cm(section.bottom_margin)
        profile["page"]["margin_left_cm"] = _safe_cm(section.left_margin)
        profile["page"]["header_distance_cm"] = _safe_cm(section.header_distance)
        profile["page"]["footer_distance_cm"] = _safe_cm(section.footer_distance)

        header_text = "\n".join((pp.text or "").strip() for pp in section.header.paragraphs if (pp.text or "").strip())
        footer_text = "\n".join((pp.text or "").strip() for pp in section.footer.paragraphs if (pp.text or "").strip())
        profile["header_footer"]["header_text"] = header_text
        profile["header_footer"]["footer_text"] = footer_text

    found = set()
    for raw in _iter_doc_strings(doc):
        for match in PLACEHOLDER_RE.finditer(raw):
            found.add(match.group(1).strip().upper())
    profile["placeholders"] = sorted(found)
    return profile



def _is_heading_paragraph(paragraph: Any) -> bool:
    style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").strip().lower()
    if style_name.startswith("heading"):
        return True
    # У стилей заголовков в документе уровень также хранится во внутренней разметке.
    ppr = getattr(getattr(paragraph, "_p", None), "pPr", None)
    return bool(ppr is not None and getattr(ppr, "outlineLvl", None) is not None)


def _is_toc_paragraph(paragraph: Any) -> bool:
    style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").strip().lower()
    return style_name.startswith("toc")


def _normalize_section_title(text: str) -> str:
    value = (text or "").strip()
    return re.sub(r"\s+", " ", value)


def _normalized_heading_key(value: str) -> str:
    clean = _normalize_section_title(value).casefold()
    clean = HEADING_PREFIX_RE.sub("", clean).strip()
    clean = re.sub(r"[^\w\s]", " ", clean)
    clean = re.sub(r"\s+", " ", clean).strip()
    return clean


def _is_special_heading_text(raw_text: str) -> bool:
    clean = _normalized_heading_key(raw_text).strip(" .:-")
    if not clean:
        return False
    return clean in SPECIAL_SECTION_TITLES or clean.startswith("\u043f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u0435")


def _is_toc_line(raw_text: str) -> bool:
    value = _normalize_section_title(raw_text)
    if not value:
        return False
    if value.isdigit():
        return True
    if TOC_DOTS_RE.search(value):
        return True
    if "\t" in value and TOC_PAGE_RE.match(value):
        return True

    return False


def _is_numbered_heading_text(raw_text: str) -> bool:
    value = _normalize_section_title(raw_text)
    match = NUMBERED_HEADING_RE.match(value)
    if not match:
        return False

    title_part = (match.group(2) or "").strip()
    if not title_part:
        return False

    if title_part.endswith((".", ";", ":", ",")):
        return False

    words = [w for w in re.split(r"\s+", title_part) if w]
    if len(words) > 18:
        return False

    return True


def _is_candidate_heading(paragraph: Any, raw_text: str) -> bool:
    if _is_toc_paragraph(paragraph):
        return False
    if _is_heading_paragraph(paragraph):
        return True
    if _is_toc_line(raw_text):
        return False
    if _is_numbered_heading_text(raw_text):
        return True
    return _is_special_heading_text(raw_text)


def _get_style_heading_level(paragraph: Any) -> int | None:
    style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").strip()
    match = HEADING_STYLE_LEVEL_RE.match(style_name)
    if not match:
        return None
    try:
        level = int(match.group(1)) - 1
    except Exception:
        return None
    if level < 0 or level > 8:
        return None
    return level


def _get_outline_level(paragraph: Any) -> int | None:
    ppr = getattr(getattr(paragraph, "_p", None), "pPr", None)
    if ppr is None:
        return None
    outline = getattr(ppr, "outlineLvl", None)
    if outline is None:
        return None
    value = getattr(outline, "val", None)
    try:
        level = int(value)
    except Exception:
        return None
    if level < 0 or level > 8:
        return None
    return level


def _next_outline_number(counters: list[int], level: int) -> str:
    for index in range(level):
        if counters[index] == 0:
            counters[index] = 1
    counters[level] += 1
    for index in range(level + 1, len(counters)):
        counters[index] = 0
    return ".".join(str(part) for part in counters[: level + 1])


def _normalized_dedupe_key(value: str) -> str:
    clean = _normalize_section_title(value).casefold()
    clean = re.sub(r"[^\w\s]", " ", clean)
    clean = re.sub(r"\s+", " ", clean).strip()
    return clean


def _apply_explicit_number_to_counters(counters: list[int], raw_text: str) -> bool:
    value = _normalize_section_title(raw_text)
    match = NUMBERED_HEADING_RE.match(value)
    if not match:
        return False

    number_part = (match.group(1) or "").strip()
    if not number_part:
        return False

    try:
        numbers = [int(part) for part in number_part.split(".") if part]
    except Exception:
        return False
    if not numbers:
        return False

    level = len(numbers) - 1
    if level >= len(counters):
        return False

    for idx, number in enumerate(numbers):
        counters[idx] = max(0, number)
    for idx in range(len(numbers), len(counters)):
        counters[idx] = 0

    return True


def _unique_section_code(template: ProjectTemplate, raw_title: str, index: int) -> str:
    base = slugify(raw_title or "")[:90] or f"section-{index}"
    candidate = base
    suffix = 2
    while ProjectTemplateSection.objects.filter(template=template, code=candidate).exists():
        suffix_str = f"-{suffix}"
        candidate = f"{base[:120 - len(suffix_str)]}{suffix_str}"
        suffix += 1
    return candidate



def _clean_heading_title_for_output(raw_text: str) -> str:
    value = _normalize_section_title(raw_text)
    if not value:
        return ""

    # Сохраняем нумерацию в начале (1 / 1.1 / 1.2.3) и убираем номер страницы в конце, если он попал в строку.
    match = NUMBERED_HEADING_RE.match(value)
    if match:
        num = (match.group(1) or "").strip()
        title = (match.group(2) or "").strip()
        title = re.sub(r"\s+\d{1,3}$", "", title).strip()
        if title:
            return f"{num} {title}"
        return num

    return value


def _extract_heading_titles_from_doc(doc: Any) -> list[str]:
    titles: list[str] = []
    seen: set[str] = set()
    outline_counters = [0] * 9

    for paragraph in doc.paragraphs:
        raw = (paragraph.text or "").strip()
        if not raw:
            continue
        if not _is_candidate_heading(paragraph, raw):
            continue

        normalized = _clean_heading_title_for_output(raw)
        if not normalized:
            continue

        # Если в шаблоне уже есть номера заголовков, считаем их источником истины.
        has_explicit_number = _apply_explicit_number_to_counters(outline_counters, normalized)
        if not has_explicit_number and _is_heading_paragraph(paragraph) and not _is_special_heading_text(normalized):
            level = _get_style_heading_level(paragraph)
            if level is None:
                level = _get_outline_level(paragraph)
            if level is not None:
                # Иначе генерируем стабильную нумерацию по уровням структуры заголовков.
                generated_number = _next_outline_number(outline_counters, level)
                normalized = f"{generated_number} {normalized}".strip()

        dedupe_key = _normalized_dedupe_key(normalized)
        if not dedupe_key or dedupe_key in seen:
            continue

        seen.add(dedupe_key)
        titles.append(normalized)
        if APPENDIX_HEADING_RE.match(normalized):
            break

    return titles


def preview_template_sections_from_docx_bytes(file_bytes: bytes) -> list[str]:
    """Предпросмотр заголовков из загруженного .docx без сохранения разделов шаблона."""
    if not file_bytes:
        return []
    try:
        from docx import Document
    except Exception:
        return []
    try:
        doc = Document(BytesIO(file_bytes))
    except Exception:
        return []
    return _extract_heading_titles_from_doc(doc)

def extract_template_sections_from_docx(template: ProjectTemplate, *, overwrite: bool = True) -> int:
    """Извлекает разделы шаблона из заголовков .docx файла."""
    if not template.template_file:
        return 0

    try:
        from docx import Document
    except Exception:
        return 0

    try:
        doc = Document(template.template_file.path)
    except Exception:
        return 0

    titles = _extract_heading_titles_from_doc(doc)

    template.format_profile = extract_template_format_profile_from_doc(doc)
    template.save(update_fields=["format_profile"])

    if not titles:
        return 0

    # При повторном импорте полностью пересобираем разделы под текущий файл шаблона.
    if overwrite:
        template.sections.all().delete()

    created = 0
    for order, title in enumerate(titles, start=1):
        ProjectTemplateSection.objects.create(
            template=template,
            title=title,
            code=_unique_section_code(template, title, order),
            order=order,
            default_task="",
        )
        created += 1
    return created


def build_template_editor_profile(template: ProjectTemplate) -> dict[str, Any]:
    if template.format_profile:
        return template.format_profile

    profile: dict[str, Any] = {
        "defaults": {"font_family": "", "font_size_pt": None},
        "paragraph": {
            "line_spacing": None,
            "space_before_pt": None,
            "space_after_pt": None,
            "first_line_indent_cm": None,
            "left_indent_cm": None,
            "right_indent_cm": None,
            "alignment": "left",
        },
        "page": {
            "margin_top_cm": None,
            "margin_right_cm": None,
            "margin_bottom_cm": None,
            "margin_left_cm": None,
            "header_distance_cm": None,
            "footer_distance_cm": None,
        },
        "header_footer": {"header_text": "", "footer_text": ""},
        "placeholders": [],
    }

    if not template.template_file:
        return profile

    try:
        from docx import Document
    except Exception:
        return profile

    try:
        doc = Document(template.template_file.path)
    except Exception:
        return profile

    return extract_template_format_profile_from_doc(doc)










