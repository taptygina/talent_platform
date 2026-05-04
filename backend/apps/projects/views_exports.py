from io import BytesIO
from copy import deepcopy
import base64
import binascii
from dataclasses import dataclass
import html
from html.parser import HTMLParser
import re

from django.http import HttpResponse
from rest_framework import permissions, status
from rest_framework.response import Response
from rest_framework.views import APIView

from apps.projects.template_utils import render_placeholders

from apps.projects.models import (
    Project,
    ProjectStageSubmission,
    ProjectStatus,
    StageStatus,
    StageSubmissionStatus,
)
from apps.users.models import UserRole


def _normalize_heading_key(text: str) -> str:
    value = re.sub(r"\s+", " ", (text or "").replace("\t", " ").strip())
    return value.lower()


def _heading_base_key(text: str) -> str:
    raw = _normalize_heading_key(text)
    raw = re.sub(r"^\d+(?:\.\d+)*\.?\)?\s*", "", raw)
    return raw.strip()


def _is_probable_heading_paragraph(paragraph) -> bool:
    raw = (getattr(paragraph, "text", "") or "").strip()
    if not raw:
        return False
    style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").strip().lower()
    if style_name.startswith("toc"):
        return False
    if "heading" in style_name or "заголов" in style_name:
        return True
    return bool(re.match(r"^\d+(?:\.\d+)*\s+\S", raw))


def _pick_template_samples(template_doc):
    heading_sample = None
    body_sample = None

    for paragraph in template_doc.paragraphs:
        raw = (paragraph.text or "").strip()
        if not raw:
            continue
        style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").strip().lower()
        if style_name.startswith("toc"):
            continue
        if heading_sample is None and _is_probable_heading_paragraph(paragraph):
            heading_sample = paragraph
            continue
        if body_sample is None and not _is_probable_heading_paragraph(paragraph):
            body_sample = paragraph
        if heading_sample is not None and body_sample is not None:
            break

    if body_sample is None:
        for paragraph in template_doc.paragraphs:
            if (paragraph.text or "").strip():
                body_sample = paragraph
                break

    if heading_sample is None:
        for paragraph in template_doc.paragraphs:
            raw = (paragraph.text or "").strip()
            if not raw:
                continue
            style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").strip().lower()
            if "heading" in style_name or "заголов" in style_name:
                heading_sample = paragraph
                break

    if heading_sample is None:
        heading_sample = body_sample

    return heading_sample, body_sample


def _build_heading_maps(doc):
    exact_map = {}
    base_map = {}

    for paragraph in doc.paragraphs:
        text_value = (paragraph.text or "").strip()
        if not text_value:
            continue

        exact = _normalize_heading_key(text_value)
        base = _heading_base_key(text_value)

        if exact and exact not in exact_map:
            exact_map[exact] = paragraph
        if base and base not in base_map:
            base_map[base] = paragraph

    return exact_map, base_map


def _build_heading_samples_map(doc):
    exact_map = {}
    base_map = {}

    for paragraph in doc.paragraphs:
        text_value = (paragraph.text or "").strip()
        if not text_value:
            continue
        if not _is_probable_heading_paragraph(paragraph):
            continue

        exact = _normalize_heading_key(text_value)
        base = _heading_base_key(text_value)

        if exact and exact not in exact_map:
            exact_map[exact] = paragraph
        if base and base not in base_map:
            base_map[base] = paragraph

    return exact_map, base_map


def _copy_paragraph_props(dst_paragraph, src_paragraph):
    if src_paragraph is None:
        return

    try:
        src_style = getattr(src_paragraph, "style", None)
        if src_style is not None and getattr(src_style, "name", None):
            dst_paragraph.style = src_style.name
    except Exception:
        pass

    try:
        dst_p = dst_paragraph._p
        src_p = src_paragraph._p
        dst_ppr = dst_p.pPr
        if dst_ppr is not None:
            dst_p.remove(dst_ppr)
        src_ppr = src_p.pPr
        if src_ppr is not None:
            dst_p.insert(0, deepcopy(src_ppr))
    except Exception:
        pass

    try:
        src_run = next((r for r in src_paragraph.runs if (r.text or "").strip()), None)
        if src_run is None:
            return
        if not dst_paragraph.runs:
            return
        dst_run = dst_paragraph.runs[0]
        src_rpr = src_run._r.rPr
        if src_rpr is None:
            return
        dst_r = dst_run._r
        dst_rpr = dst_r.rPr
        if dst_rpr is not None:
            dst_r.remove(dst_rpr)
        dst_r.insert(0, deepcopy(src_rpr))
    except Exception:
        pass


def _append_cloned_heading_paragraph(doc, src_paragraph, fallback_text: str):
    if src_paragraph is None:
        paragraph = doc.add_paragraph((fallback_text or "").strip())
        return paragraph

    try:
        body = doc._element.body
        cloned = deepcopy(src_paragraph._p)
        sect_idx = None
        for idx, child in enumerate(list(body)):
            if child.tag.endswith("}sectPr"):
                sect_idx = idx
                break
        if sect_idx is None:
            body.append(cloned)
        else:
            body.insert(sect_idx, cloned)

        from docx.text.paragraph import Paragraph

        new_par = Paragraph(cloned, doc)
        text_value = re.sub(r"\s+", " ", (fallback_text or "").replace("\t", " ")).strip()
        if text_value:
            if new_par.runs:
                new_par.runs[0].text = text_value
                for run in new_par.runs[1:]:
                    run.text = ""
            else:
                new_par.add_run(text_value)
        return new_par
    except Exception:
        paragraph = doc.add_paragraph((fallback_text or "").strip())
        _copy_paragraph_props(paragraph, src_paragraph)
        return paragraph


def _insert_template_paragraph_after(anchor_paragraph, text_value: str, sample_paragraph):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    anchor_paragraph._p.addnext(new_p)
    new_par = Paragraph(new_p, anchor_paragraph._parent)
    _copy_paragraph_props(new_par, sample_paragraph)
    clean_text = re.sub(r"\s+", " ", (text_value or "").replace("\t", " ")).strip()
    if clean_text:
        new_par.add_run(clean_text)
    return new_par


def _clear_section_body_between_headings(current_heading, next_heading):
    parent = current_heading._p.getparent()
    node = current_heading._p.getnext()
    stop_node = next_heading._p if next_heading is not None else None
    while node is not None and node is not stop_node:
        next_node = node.getnext()
        parent.remove(node)
        node = next_node


@dataclass
class _StyledRun:
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False


@dataclass
class _ContentBlock:
    kind: str
    runs: list[_StyledRun]
    image_src: str = ""


class _SubmissionHtmlParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocks: list[_ContentBlock] = []
        self._style_stack: list[dict[str, bool]] = []
        self._in_list = False
        self._current_block: _ContentBlock | None = None

    def handle_starttag(self, tag, attrs):
        name = (tag or "").lower()
        attr_map = {key.lower(): value for key, value in (attrs or []) if key}
        if name == "br":
            self._append_text("\n")
            return
        if name in {"p", "div"}:
            self._start_block("paragraph")
            return
        if name in {"ul", "ol"}:
            self._in_list = True
            return
        if name == "li":
            self._start_block("list_item")
            return
        if name in {"strong", "b", "em", "i", "u"}:
            self._style_stack.append(
                {
                    "bold": name in {"strong", "b"},
                    "italic": name in {"em", "i"},
                    "underline": name == "u",
                }
            )
            return
        if name == "a":
            self._style_stack.append({"bold": False, "italic": False, "underline": True})
            return
        if name == "img":
            src = (attr_map.get("src") or "").strip()
            self._close_block()
            if src:
                self.blocks.append(_ContentBlock(kind="image", runs=[], image_src=src))

    def handle_endtag(self, tag):
        name = (tag or "").lower()
        if name in {"p", "div", "li"}:
            self._close_block()
            return
        if name in {"ul", "ol"}:
            self._in_list = False
            return
        if name in {"strong", "b", "em", "i", "u", "a"} and self._style_stack:
            self._style_stack.pop()

    def handle_data(self, data):
        self._append_text(data or "")

    def close(self):
        super().close()
        self._close_block()

    def _current_styles(self) -> dict[str, bool]:
        styles = {"bold": False, "italic": False, "underline": False}
        for item in self._style_stack:
            styles["bold"] = styles["bold"] or item.get("bold", False)
            styles["italic"] = styles["italic"] or item.get("italic", False)
            styles["underline"] = styles["underline"] or item.get("underline", False)
        return styles

    def _start_block(self, kind: str):
        if self._current_block is None:
            self._current_block = _ContentBlock(kind=kind, runs=[])
            return
        if self._current_block.runs:
            self._close_block()
        self._current_block = _ContentBlock(kind=kind, runs=[])

    def _close_block(self):
        if self._current_block is None:
            return
        if self._current_block.kind == "image":
            if self._current_block.image_src:
                self.blocks.append(self._current_block)
        elif any(run.text.strip() for run in self._current_block.runs):
            self.blocks.append(self._current_block)
        self._current_block = None

    def _append_text(self, value: str):
        if not value:
            return
        text_value = value.replace("\r\n", "\n").replace("\r", "\n")
        chunks = text_value.split("\n")
        for idx, chunk in enumerate(chunks):
            normalized = re.sub(r"[ \t]+", " ", chunk)
            if normalized:
                block_kind = "list_item" if self._in_list else "paragraph"
                if self._current_block is None:
                    self._current_block = _ContentBlock(kind=block_kind, runs=[])
                styles = self._current_styles()
                self._current_block.runs.append(
                    _StyledRun(
                        text=normalized,
                        bold=styles["bold"],
                        italic=styles["italic"],
                        underline=styles["underline"],
                    )
                )
            if idx != len(chunks) - 1:
                self._close_block()


def _extract_blocks_from_submission(value: str) -> list[_ContentBlock]:
    source = (value or "").strip()
    if not source:
        return []

    if "<" not in source and ">" not in source:
        plain_text = html.unescape(source)
        plain_text = plain_text.replace("\r\n", "\n").replace("\r", "\n")
        blocks = []
        for line in plain_text.split("\n"):
            normalized = re.sub(r"[ \t]+", " ", line).strip()
            if not normalized:
                continue
            blocks.append(_ContentBlock(kind="paragraph", runs=[_StyledRun(text=normalized)]))
        return blocks

    parser = _SubmissionHtmlParser()
    parser.feed(source)
    parser.close()
    return parser.blocks


def _copy_run_props(dst_run, src_run):
    if src_run is None:
        return
    try:
        src_rpr = src_run._r.rPr
        if src_rpr is None:
            return
        dst_r = dst_run._r
        dst_rpr = dst_r.rPr
        if dst_rpr is not None:
            dst_r.remove(dst_rpr)
        dst_r.insert(0, deepcopy(src_rpr))
    except Exception:
        return


DATA_URL_IMAGE_RE = re.compile(r"^data:image/([a-zA-Z0-9.+-]+);base64,(.+)$", re.DOTALL)
DOCX_NATIVE_IMAGE_TYPES = {"png", "jpg", "jpeg", "gif", "bmp", "tif", "tiff"}


def _decode_data_url_image(image_src: str) -> tuple[bytes, str] | tuple[None, None]:
    source = (image_src or "").strip()
    if not source:
        return None, None
    match = DATA_URL_IMAGE_RE.match(source)
    if not match:
        return None, None

    image_type = (match.group(1) or "").lower()
    encoded = (match.group(2) or "").strip()
    encoded = re.sub(r"\s+", "", encoded)
    if not encoded:
        return None, None

    missing_padding = len(encoded) % 4
    if missing_padding:
        encoded += "=" * (4 - missing_padding)

    try:
        image_bytes = base64.b64decode(encoded, validate=False)
    except (binascii.Error, ValueError):
        return None, None
    if not image_bytes:
        return None, None
    return image_bytes, image_type


def _prepare_docx_image_bytes(image_bytes: bytes, image_type: str) -> bytes | None:
    if image_type in DOCX_NATIVE_IMAGE_TYPES:
        return image_bytes

    try:
        from PIL import Image
    except Exception:
        return None

    try:
        with Image.open(BytesIO(image_bytes)) as image:
            converted = BytesIO()
            if image.mode in {"RGBA", "LA", "P"}:
                image = image.convert("RGBA")
            else:
                image = image.convert("RGB")
            image.save(converted, format="PNG")
            return converted.getvalue()
    except Exception:
        return None


def _fill_paragraph_from_block(paragraph, block: _ContentBlock, sample_paragraph):
    if block.kind == "image":
        image_bytes, image_type = _decode_data_url_image(block.image_src)
        if not image_bytes or not image_type:
            paragraph.text = "[изображение не поддерживается]"
            _copy_paragraph_props(paragraph, sample_paragraph)
            return
        prepared_bytes = _prepare_docx_image_bytes(image_bytes, image_type)
        if not prepared_bytes:
            paragraph.text = "[изображение не поддерживается]"
            _copy_paragraph_props(paragraph, sample_paragraph)
            return
        try:
            from docx.shared import Cm

            run = paragraph.add_run()
            run.add_picture(BytesIO(prepared_bytes), width=Cm(14.5))
        except Exception:
            paragraph.text = "[ошибка вставки изображения]"
            _copy_paragraph_props(paragraph, sample_paragraph)
        return

    if block.kind == "list_item":
        raw_text = "".join(run.text for run in block.runs).strip()
        paragraph.text = f"- {raw_text}" if raw_text else "-"
        _copy_paragraph_props(paragraph, sample_paragraph)
        return

    for run in list(paragraph.runs):
        run.text = ""

    base_run_sample = next((run for run in getattr(sample_paragraph, "runs", []) if (run.text or "").strip()), None)
    for item in block.runs:
        piece = item.text
        if not piece:
            continue
        run = paragraph.add_run(piece)
        _copy_run_props(run, base_run_sample)
        if item.bold:
            run.bold = True
        if item.italic:
            run.italic = True
        if item.underline:
            run.underline = True


class ProjectExportNirsDocxView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, pk):
        try:
            project = Project.objects.select_related("supervisor", "template").get(pk=pk)
        except Project.DoesNotExist:
            return Response({"detail": "Проект не найден."}, status=status.HTTP_404_NOT_FOUND)

        user = request.user
        if user.role not in {UserRole.TEACHER, UserRole.CURATOR, UserRole.ADMIN}:
            return Response({"detail": "Недостаточно прав."}, status=status.HTTP_403_FORBIDDEN)
        if user.role == UserRole.TEACHER and project.supervisor_id != user.id:
            return Response({"detail": "Недостаточно прав."}, status=status.HTTP_403_FORBIDDEN)
        if project.status != ProjectStatus.DONE:
            return Response({"detail": "Проект должен быть в статусе «Завершен»."}, status=status.HTTP_400_BAD_REQUEST)
        if project.stages.exclude(status=StageStatus.APPROVED).exists():
            return Response({"detail": "Перед выгрузкой все этапы должны быть в статусе «Принят»."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            from docx import Document
        except Exception:
            return Response(
                {"detail": "Для генерации .docx установите зависимость python-docx."},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        if not project.template_id or not project.template:
            return Response(
                {"detail": "Шаблон не найден в проекте .docx. Привяжите шаблон к редактируемому проекту."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        if not getattr(project.template, "template_file", None):
            return Response(
                {"detail": "У выбранного шаблона не загружен .docx файл. Загрузите файл шаблона и повторите экспорт."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            doc = Document(project.template.template_file.path)
        except Exception:
            return Response(
                {"detail": "Не удалось открыть файл шаблона .docx. Проверьте файл и повторите попытку."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        heading_sample, body_sample = _pick_template_samples(doc)
        sample_exact_map, sample_base_map = _build_heading_samples_map(doc)
        target_exact_map, target_base_map = _build_heading_maps(doc)

        all_stages = list(project.stages.order_by("order", "id"))
        stage_by_template_section = {
            stage.template_section_id: stage
            for stage in all_stages
            if stage.template_section_id
        }
        stage_by_order = {}
        for stage in all_stages:
            stage_by_order.setdefault(stage.order, stage)

        sections = []
        if project.template.sections.exists():
            for section in project.template.sections.order_by("order", "id"):
                stage = stage_by_template_section.get(section.id) or stage_by_order.get(section.order)
                sections.append((section.title, stage))
        else:
            for stage in all_stages:
                sections.append((stage.title, stage))

        approved_submissions_by_stage = {}
        approved_submissions = (
            ProjectStageSubmission.objects.select_related("student", "stage")
            .filter(stage__project=project, status=StageSubmissionStatus.APPROVED)
            .order_by("stage__order", "stage__id", "student__last_name", "student__first_name")
        )
        for submission in approved_submissions:
            approved_submissions_by_stage.setdefault(submission.stage_id, []).append(submission)

        section_rows = []
        for section_title, stage in sections:
            exact_key = _normalize_heading_key(section_title)
            base_key = _heading_base_key(section_title)

            sample_heading = sample_exact_map.get(exact_key) or sample_base_map.get(base_key) or heading_sample
            target_heading = target_exact_map.get(exact_key) or target_base_map.get(base_key)

            if target_heading is None:
                target_heading = _append_cloned_heading_paragraph(doc, sample_heading, section_title)
                target_exact_map[exact_key] = target_heading
                if base_key and base_key not in target_base_map:
                    target_base_map[base_key] = target_heading

            section_rows.append({"title": section_title, "stage": stage, "heading": target_heading})

        # Важно: перед вставкой новых данных очищаем старое содержимое каждого раздела.
        for idx, row in enumerate(section_rows):
            next_heading = section_rows[idx + 1]["heading"] if idx + 1 < len(section_rows) else None
            _clear_section_body_between_headings(row["heading"], next_heading)

        for row in section_rows:
            stage = row["stage"]
            if not stage:
                continue

            anchor = row["heading"]
            for submission in approved_submissions_by_stage.get(stage.id, []):
                submission_text = render_placeholders(
                    submission.submission_text or "",
                    student=None,
                    supervisor=project.supervisor,
                    project=project,
                )
                blocks = _extract_blocks_from_submission(submission_text)
                if not blocks:
                    continue
                for block in blocks:
                    anchor = _insert_template_paragraph_after(anchor, "", body_sample)
                    _fill_paragraph_from_block(anchor, block, body_sample)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"nirs_project_{project.id}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response


class ProjectExportMatrixXlsxView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, pk):
        try:
            project = Project.objects.select_related("supervisor").prefetch_related("participants", "stages").get(pk=pk)
        except Project.DoesNotExist:
            return Response({"detail": "Проект не найден."}, status=status.HTTP_404_NOT_FOUND)

        user = request.user
        if user.role not in {UserRole.TEACHER, UserRole.CURATOR, UserRole.ADMIN, UserRole.METHODIST}:
            return Response({"detail": "Недостаточно прав."}, status=status.HTTP_403_FORBIDDEN)
        if user.role == UserRole.TEACHER and project.supervisor_id != user.id:
            return Response({"detail": "Недостаточно прав."}, status=status.HTTP_403_FORBIDDEN)

        try:
            from openpyxl import Workbook
            from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
            from openpyxl.utils import get_column_letter
        except Exception:
            return Response(
                {"detail": "Для генерации .xlsx установите зависимость openpyxl."},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        status_label = {
            "draft": "Не сдал",
            "submitted": "На проверке",
            "needs_changes": "Не сдал",
            "approved": "Сдал",
            "none": "Не сдал",
        }
        project_status_label = {
            "planned": "Запланирован",
            "in_progress": "В работе",
            "review": "На проверке",
            "done": "Завершен",
            "cancelled": "Отменен",
        }
        status_fill = {
            "draft": "F4CCCC",
            "submitted": "FFF2CC",
            "needs_changes": "F4CCCC",
            "approved": "D9EAD3",
            "none": "F4CCCC",
        }
        thin = Side(style="thin", color="000000")
        table_border = Border(left=thin, right=thin, top=thin, bottom=thin)
        font_regular = Font(name="Times New Roman", size=12)
        font_header = Font(name="Times New Roman", size=12, bold=True)
        font_title = Font(name="Times New Roman", size=14, bold=True)

        students = list(project.participants.filter(role=UserRole.STUDENT).order_by("last_name", "first_name"))
        stages = list(project.stages.order_by("order", "id"))

        wb = Workbook()
        ws = wb.active
        ws.title = "Матрица этапов"

        ws["A1"] = "Матрица этапов проекта"
        ws["A1"].font = font_title
        ws["A2"] = f"Проект: {project.title}"
        ws["A2"].font = font_regular
        ws["A3"] = f"Руководитель: {project.supervisor.full_name or project.supervisor.username}"
        ws["A3"].font = font_regular
        ws["A4"] = f"Статус проекта: {project_status_label.get(project.status, project.status)}"
        ws["A4"].font = font_regular

        header_row = 6
        ws.cell(row=header_row, column=1, value="Студент")
        ws.cell(row=header_row, column=1).font = font_header
        ws.cell(row=header_row, column=1).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.cell(row=header_row, column=1).border = table_border

        for col_idx, stage in enumerate(stages, start=2):
            ws.cell(row=header_row, column=col_idx, value=f"{stage.order}. {stage.title}")
            ws.cell(row=header_row, column=col_idx).font = font_header
            ws.cell(row=header_row, column=col_idx).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            ws.cell(row=header_row, column=col_idx).border = table_border

        for row_idx, student in enumerate(students, start=header_row + 1):
            student_name = student.full_name or student.username
            ws.cell(row=row_idx, column=1, value=student_name)
            ws.cell(row=row_idx, column=1).font = font_regular
            ws.cell(row=row_idx, column=1).border = table_border
            ws.cell(row=row_idx, column=1).alignment = Alignment(vertical="center", wrap_text=True)
            for col_idx, stage in enumerate(stages, start=2):
                submission = (
                    ProjectStageSubmission.objects.filter(stage=stage, student=student)
                    .order_by("-updated_at")
                    .first()
                )
                if submission:
                    key = submission.status
                    label = status_label.get(key, submission.status)
                else:
                    key = "none"
                    label = status_label[key]
                cell = ws.cell(row=row_idx, column=col_idx, value=label)
                cell.font = font_regular
                cell.alignment = Alignment(horizontal="center", vertical="center")
                color = status_fill.get(key, "F4CCCC")
                cell.fill = PatternFill(start_color=color, end_color=color, fill_type="solid")
                cell.border = table_border

        ws.column_dimensions["A"].width = 34
        for col_idx in range(2, len(stages) + 2):
            ws.column_dimensions[get_column_letter(col_idx)].width = 22
        ws.row_dimensions[header_row].height = 42
        for row_idx in range(header_row + 1, header_row + 1 + len(students)):
            ws.row_dimensions[row_idx].height = 28
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(2, len(stages) + 1))

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        filename = f"project_matrix_{project.id}.xlsx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response
